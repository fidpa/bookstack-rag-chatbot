"""
Indexing Service für Knowledge Base
Verwaltet die Volltext-Indexierung der Dokumente
"""

import os
import logging
from datetime import datetime
from typing import Optional, Tuple, List, Dict

from utils.database import get_db_connection
from ..models import KnowledgeDocument
from .chunking import ChunkingService, DocumentChunk

logger = logging.getLogger(__name__)


def _extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract plain text from a file based on its MIME type or extension."""
    ft = file_type.lower()

    # PDF
    if 'pdf' in ft or file_path.lower().endswith('.pdf'):
        try:
            import pypdfium2 as pdfium
            doc = pdfium.PdfDocument(file_path)
            pages = [doc[i].get_textpage().get_text_range() for i in range(len(doc))]
            return '\n\n'.join(pages)
        except Exception:
            pass
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            return '\n\n'.join(p.extract_text() or '' for p in reader.pages)
        except Exception as e:
            raise RuntimeError(f"PDF extraction failed: {e}") from e

    # DOCX
    if 'wordprocessingml' in ft or 'docx' in ft or file_path.lower().endswith('.docx'):
        try:
            from docx import Document
            doc = Document(file_path)
            return '\n'.join(p.text for p in doc.paragraphs)
        except Exception as e:
            raise RuntimeError(f"DOCX extraction failed: {e}") from e

    # Plain text / Markdown
    if 'text' in ft or file_path.lower().endswith(('.md', '.txt', '.markdown')):
        with open(file_path, encoding='utf-8', errors='replace') as fh:
            return fh.read()

    raise ValueError(f"Unsupported file type: {file_type}")


class IndexingService:
    """Service für die Volltext-Indexierung von Wissensbasis-Dokumenten"""

    # Maximale Chunk-Größe für Indexierung (10KB)
    MAX_CHUNK_SIZE = 10240

    @classmethod
    def extract_text_from_document(cls, doc: KnowledgeDocument) -> Tuple[bool, str]:
        """
        Extrahiert Text aus einem Dokument.

        Returns:
            (success, text_content)
        """
        try:
            if not os.path.exists(doc.file_path):
                logger.error(f"File not found: {doc.file_path}")
                return False, f"File not found: {doc.file_path}"

            logger.info(f"Extracting text from: {doc.file_path}")
            text = _extract_text_from_file(doc.file_path, doc.file_type)

            if not text or not text.strip():
                logger.warning(f"Empty text extracted for {doc.original_filename}")
                return False, "No text extracted"

            logger.info(f"Text extraction successful for {doc.original_filename}: {len(text)} chars")
            return True, text

        except Exception as e:
            error_msg = f"Text extraction failed for {doc.original_filename}: {e}"
            logger.error(error_msg, exc_info=True)
            return False, error_msg
    
    @classmethod
    def index_document(cls, doc_id: int, use_chunking: bool = True) -> Tuple[bool, str]:
        """
        Indexiert ein einzelnes Dokument für die Volltext-Suche
        
        Args:
            doc_id: ID des zu indexierenden Dokuments
            use_chunking: Ob das neue Chunking-System verwendet werden soll
        
        Returns:
            (success, message)
        """
        try:
            with get_db_connection() as conn:
                cursor = conn.cursor()
                
                # Dokument laden
                cursor.execute('SELECT * FROM kb_documents WHERE id = ?', (doc_id,))
                row = cursor.fetchone()
                
                if not row:
                    return False, "Dokument nicht gefunden"
                
                doc = KnowledgeDocument.from_db_row(dict(row))
                
                # Text extrahieren
                success, text_content = cls.extract_text_from_document(doc)
                
                if not success:
                    logger.error(f"Text extraction failed for doc {doc_id}: {text_content}")
                    return False, f"Text-Extraktion fehlgeschlagen: {text_content}"
                
                if use_chunking:
                    return cls._index_with_chunking(doc_id, text_content, doc)
                else:
                    return cls._index_legacy(doc_id, text_content, doc)

        except Exception as e:
            logger.error(f"Error indexing document {doc_id}: {e}")
            return False, f"Indexing failed: {e}"

    @classmethod
    def _index_with_chunking(cls, doc_id: int, text_content: str, doc: KnowledgeDocument) -> Tuple[bool, str]:
        """Index document using the chunking pipeline."""
        try:
            with get_db_connection() as conn:
                cursor = conn.cursor()

                cursor.execute('DELETE FROM kb_chunks WHERE doc_id = ?', (doc_id,))

                chunking_service = ChunkingService()
                chunks = chunking_service.chunk_document(text_content, doc_id)

                if not chunks:
                    logger.warning(f"No chunks created for document {doc_id}")
                    return False, "No chunks created"

                for chunk in chunks:
                    chunk_dict = chunk.to_dict()
                    cursor.execute('''
                        INSERT INTO kb_chunks
                        (doc_id, chunk_index, chunk_text, start_pos, end_pos, word_count)
                        VALUES (?, ?, ?, ?, ?, ?)
                    ''', (
                        chunk_dict['doc_id'],
                        chunk_dict['chunk_index'],
                        chunk_dict['chunk_text'],
                        chunk_dict['start_pos'],
                        chunk_dict['end_pos'],
                        chunk.word_count
                    ))

                stats = chunking_service.get_chunk_statistics(chunks)
                cursor.execute('''
                    INSERT OR REPLACE INTO kb_chunk_stats
                    (doc_id, total_chunks, avg_chunk_size, min_chunk_size, max_chunk_size)
                    VALUES (?, ?, ?, ?, ?)
                ''', (
                    doc_id,
                    stats['total_chunks'],
                    int(stats['avg_words_per_chunk']),
                    stats['min_words'],
                    stats['max_words']
                ))

                cursor.execute('''
                    UPDATE kb_documents
                    SET last_indexed = ?,
                        chunking_status = 'completed',
                        chunk_count = ?
                    WHERE id = ?
                ''', (datetime.now().isoformat(), len(chunks), doc_id))

                # also update legacy FTS index for backwards compatibility
                cursor.execute('DELETE FROM kb_search_fts WHERE doc_id = ?', (doc_id,))
                cursor.execute('''
                    INSERT INTO kb_search_fts (doc_id, title, content)
                    VALUES (?, ?, ?)
                ''', (doc_id, doc.title or doc.original_filename, text_content[:cls.MAX_CHUNK_SIZE]))
                
                conn.commit()

                logger.info(f"Document {doc_id} indexed with {len(chunks)} chunks")
                return True, f"Document indexed ({len(chunks)} chunks)"

        except Exception as e:
            logger.error(f"Chunk indexing error: {e}", exc_info=True)
            return False, f"Chunk indexing failed: {e}"

    @classmethod
    def _index_legacy(cls, doc_id: int, text_content: str, doc: KnowledgeDocument) -> Tuple[bool, str]:
        """Legacy indexing without chunking (single FTS5 row per document)."""
        try:
            with get_db_connection() as conn:
                cursor = conn.cursor()

                cursor.execute('DELETE FROM kb_search_fts WHERE doc_id = ?', (doc_id,))

                slices = [text_content[i:i + cls.MAX_CHUNK_SIZE]
                          for i in range(0, len(text_content), cls.MAX_CHUNK_SIZE)]

                for slice_text in slices:
                    cursor.execute('''
                        INSERT INTO kb_search_fts (doc_id, title, content)
                        VALUES (?, ?, ?)
                    ''', (doc_id, doc.title or doc.original_filename, slice_text))

                cursor.execute('''
                    UPDATE kb_documents
                    SET last_indexed = ?,
                        chunking_status = 'legacy'
                    WHERE id = ?
                ''', (datetime.now().isoformat(), doc_id))

                conn.commit()

                logger.info(f"Document {doc_id} indexed (legacy)")
                return True, "Document indexed"

        except Exception as e:
            logger.error(f"Legacy indexing error: {e}")
            return False, f"Indexing failed: {e}"

    @classmethod
    def index_all_documents(cls, progress_callback=None) -> Tuple[int, int]:
        """
        Re-index all active documents.

        Args:
            progress_callback: Optional callback(current, total, doc_name)

        Returns:
            (successful_count, failed_count)
        """
        successful = 0
        failed = 0

        try:
            with get_db_connection() as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT id, original_filename, title
                    FROM kb_documents
                    WHERE is_active = 1
                ''')

                documents = cursor.fetchall()
                total = len(documents)

            for idx, doc in enumerate(documents):
                doc_id = doc['id']
                doc_name = doc['title'] or doc['original_filename']

                if progress_callback:
                    progress_callback(idx + 1, total, doc_name)

                success, _ = cls.index_document(doc_id)
                if success:
                    successful += 1
                else:
                    failed += 1

            logger.info(f"Re-index complete: {successful} ok, {failed} failed")

        except Exception as e:
            logger.error(f"Full re-index error: {e}")
            
        return successful, failed
    
    @classmethod
    def extract_keywords(cls, text: str, max_keywords: int = 10) -> List[str]:
        """Extract the most significant keywords from text using TF-IDF."""
        from .keyword_extraction import KeywordExtractor

        extractor = KeywordExtractor()
        keywords_with_scores = extractor.extract_keywords(
            text,
            max_keywords=max_keywords,
            include_ngrams=True
        )
        keywords = [keyword for keyword, score in keywords_with_scores]
        logger.debug(f"Extracted {len(keywords)} keywords via TF-IDF")
        return keywords

    @classmethod
    def extract_document_entities(cls, text: str) -> Dict[str, List[str]]:
        """Extract named entities from text."""
        from .keyword_extraction import KeywordExtractor

        extractor = KeywordExtractor()
        entities = extractor.extract_entities(text)
        logger.debug(f"Extracted {sum(len(v) for v in entities.values())} entities")
        return entities